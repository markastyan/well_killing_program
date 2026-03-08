from enter_initial_data import *
from enter_user_data import *
from initial_calculations import *
import base_method
import driller_method
from tooltip import Tooltip
from utility import *
from get_json import load_json_as_tuple

import numpy as np
import traceback
import logging
from docx import Document
from docx.shared import Inches
import customtkinter as ctk
from matplotlib.figure import Figure
from tkinter import messagebox
from tkinter import ttk
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from PIL import Image

initial_data = {}
initial_data_multiple_pipes = {}
user_data = {}
well_volumes = {}
boundary_parameters = {}
kick_analysis = {}
base_method_values = {}
driller_method_values = {}
multiple_pipes_values = {}

initial_data_frame = None
user_data_frame = None
kill_method_frame = None
results_frame = None
multiple_pipes_frame = None
frame_width = 700
units = {
    "Внутренний объём бурильного инструмента": "м³",
    "Объём кольцевого пространства": "м³",
    "Максимально допустимое давление на устье": "Па",
    "Эквивалентная плотность начала поглощения": "кг/м³",
    "Эквивалентная плотность начала гидроразрыва": "кг/м³",
    "Максимально допустимое давление на устье для недопущения поглощения": "Па",
    "Максимальное допускаемое давление на устье": "Па",
    "Удельный объём 1 метра кольцевого пространства инструмент – скважина": "м³/м",
    "Удельный объём 1 метра открытого ствола": "м³/м",
    "Высота столба флюида в кольцевом пространстве (считая от забоя)": "м",
    "Плотность пластового флюида": "кг/м³",
    "Тип флюида": "-",
    "Пластовое давление": "Па",
    "Плотность жидкости глушения": "кг/м³",
    "Начальное давление на стояке": "Па",
    "Скорость потока в трубах": "м/с",
    "Критерий Рейнольдса в трубах": "-",
    "Коэффициент трения": "-",
    "Скорость потока в трубе ": "м/с",
    "Критерий Рейнольдса в трубе ": "-",
    "Коэффициент трения в трубе ": "-",
    "Потери давления на трение": "Па",
    "Конечное давление на стояке": "Па",
    "Давление циркуляции на пониженной подаче": "Па",
    "Начальное давление циркуляции": "Па",
    "Фактическое пластовое давление вскрытого пласта": "Па",
    "Дополнительное приращение плотности": "кг/м³",
    "Конечное давление циркуляции при глушении скважины": "Па",
    "Градиент распределения избыточного трубного давления": "-",
    "Максимальное увеличение объема газовой пачки": "м³",
    "Максимальное давление в кольцевом пространстве": "Па",
    "Точка максимального давления на графике первой стадии глушения": "-"
}

is_multiple_pipes = False
is_extra_graph = False
is_user_graph_base = False
is_user_graph_driller = False


def create_entry(parent, placeholder_text, units, entry_width=450):
    frame = ctk.CTkFrame(parent)
    frame.pack(fill='x', pady=5)
    entry = ctk.CTkEntry(frame, placeholder_text=placeholder_text, width=entry_width)
    entry.grid(row=0, column=0, padx=((frame_width - entry_width)/2, 0), sticky='w')
    units_label = ctk.CTkLabel(frame, text=units)
    units_label.grid(row=0, column=1, padx=(5, 5), sticky='w')
    Tooltip(entry, placeholder_text)

    return entry


def initial_data_input():
    def save_initial_data():
        global initial_data
        initial_data = enter_initial_data_manual(to_float(inner_diameter_entry.get()),
                                                 to_float(section_length_entry.get()),
                                                 to_float(interval_diameter_entry.get()),
                                                 to_float(outer_diameter_entry.get()),
                                                 to_float(interval_length_entry.get()),
                                                 to_float(current_depth_entry.get()),
                                                 to_float(lowering_depth_entry.get()))
        save_values_message(initial_data)

    def initial_data_from_file():
        global initial_data
        initial_data = enter_initial_data_from_file("initial_data")
        save_values_message(initial_data)

    content_frame.pack(fill="y", padx=30, pady=10, side="right")
    ctk.CTkLabel(content_frame, text="Автоматическая загрузка").pack(fill='x', pady=5)
    from_file_button = ctk.CTkButton(content_frame, text="Загрузить данные из WellPro", command=initial_data_from_file)
    from_file_button.pack(pady=10)
    ttk.Separator(content_frame, orient='horizontal').pack(fill='x', pady=10)
    ctk.CTkLabel(content_frame, text="Ручной ввод").pack(fill='x', pady=5)
    inner_diameter_entry = create_entry(content_frame, "Внутренний диаметр одноразмерного типа труб", "м")
    section_length_entry = create_entry(content_frame, "Длина секции одноразмерного типа труб", "м")
    interval_diameter_entry = create_entry(content_frame, "Диаметр интервала", "м")
    outer_diameter_entry = create_entry(content_frame, "Наружный диаметр одноразмерного бурильного инструмента", "м")
    interval_length_entry = create_entry(content_frame, "Длина интервала", "м")
    current_depth_entry = create_entry(content_frame, "Глубина текущего забоя", "м")
    lowering_depth_entry = create_entry(content_frame, "Глубина спуска инструмента", "м")
    save_button = ctk.CTkButton(content_frame, text="Сохранить данные", command=save_initial_data)
    save_button.pack(pady=10)


def multiple_pipes_input():
    multiple_pipes_local_values = {'inner_pipe_values': {}, 'outer_pipe_values': {}}
    local_values = {}
    inner_pipe_entries = []
    outer_pipe_entries = []
    counters = [0, 0]

    def initial_data_from_file():
        global initial_data, is_multiple_pipes
        is_multiple_pipes = False
        initial_data = enter_initial_data_from_file("initial_data")
        save_values_message(initial_data)

    def save_multiple_pipes_data():
        global is_multiple_pipes, initial_data_multiple_pipes
        if to_int(outer_pipe_amount_entry.get()) == 1 and to_int(inner_pipe_amount_entry.get()) == 1:
            is_multiple_pipes = False
            global initial_data
            for i, entries in enumerate(inner_pipe_entries):
                multiple_pipes_local_values['inner_pipe_values'][i + 1] = [to_float(e.get()) for e in entries]
            for i, entries in enumerate(outer_pipe_entries):
                multiple_pipes_local_values['outer_pipe_values'][i + 1] = [to_float(e.get()) for e in entries]
            print(list(multiple_pipes_local_values['outer_pipe_values'].values()))
            print(list(multiple_pipes_local_values['inner_pipe_values'].values()))
            initial_data = enter_initial_data_manual(to_float(str(list(multiple_pipes_local_values['inner_pipe_values'].values())[0][1])),
                                                     to_float(str(list(multiple_pipes_local_values['inner_pipe_values'].values())[0][2])),
                                                     to_float(str(list(multiple_pipes_local_values['inner_pipe_values'].values())[0][0])),
                                                     to_float(str(list(multiple_pipes_local_values['outer_pipe_values'].values())[0][0])),
                                                     to_float(str(list(multiple_pipes_local_values['outer_pipe_values'].values())[0][1])),
                                                     to_float(local_values['current_depth'].get()),
                                                     to_float(local_values['lowering_depth'].get()))
            save_values_message(initial_data)
            print(initial_data)

        else:
            for i, entries in enumerate(inner_pipe_entries):
                multiple_pipes_local_values['inner_pipe_values'][i + 1] = [to_float(e.get()) for e in entries]

            for i, entries in enumerate(outer_pipe_entries):
                multiple_pipes_local_values['outer_pipe_values'][i + 1] = [to_float(e.get()) for e in entries]
            initial_data_multiple_pipes = enter_initial_data_manual_multiple_pipes(
                list(multiple_pipes_local_values['outer_pipe_values'].values()),
                list(multiple_pipes_local_values['inner_pipe_values'].values()),
                to_float(local_values['current_depth'].get()),
                to_float(local_values['lowering_depth'].get()))
            save_values_message(initial_data_multiple_pipes)
            is_multiple_pipes = True

    def add_inner_pipe():
        counters[0] += 1
        element_outer_diameter_entry = create_entry(content_frame, f"Наружный диаметр элемента {counters[0]}", "м")
        element_inner_length_entry = create_entry(content_frame, f"Внутренний диаметр элемента {counters[0]}", "м")
        element_length_entry = create_entry(content_frame, f"Длина секции элемента {counters[0]}", "м")
        inner_pipe_entries.append([element_outer_diameter_entry, element_inner_length_entry, element_length_entry])

    def add_outer_pipe():
        counters[1] += 1
        interval_diameter_entry = create_entry(content_frame, f"Диаметр интервала {counters[1]}", "м")
        interval_length_entry = create_entry(content_frame, f"Длина интервала {counters[1]}", "м")
        outer_pipe_entries.append([interval_diameter_entry, interval_length_entry])

    def add_pipe_fields():
        outer_pipe_value = outer_pipe_amount_entry.get()
        inner_pipe_value = inner_pipe_amount_entry.get()
        outer_pipe_amount_entry.pack_forget()
        inner_pipe_amount_entry.pack_forget()
        add_fields_button.pack_forget()
        label.pack_forget()
        ctk.CTkLabel(content_frame, text="Введите данные для глубины").pack(fill='x', pady=5)
        local_values['current_depth'] = create_entry(content_frame, "Глубина текущего забоя", "м")
        local_values['lowering_depth'] = create_entry(content_frame,
                                                      "Глубина спуска инструмента", "м")
        ctk.CTkLabel(content_frame, text="Введите данные для элементов инструмента").pack(fill='x', pady=5)
        for i in range(to_int(outer_pipe_value)):
            add_inner_pipe()
        ctk.CTkLabel(content_frame, text="Введите данные для интервалов скважины").pack(fill='x', pady=5)
        for i in range(to_int(inner_pipe_value)):
            add_outer_pipe()
        save_values_button = ctk.CTkButton(content_frame, text="Сохранить данные", command=save_multiple_pipes_data)
        save_values_button.pack(pady=10)

    def initial_data_from_file_multiple_pipes():
        global initial_data_multiple_pipes, is_multiple_pipes
        initial_data_multiple_pipes = enter_initial_data_from_file_multiple_pipes("initial_data_multiple_pipes")
        save_values_message(initial_data_multiple_pipes)
        outer_pipe_amount_entry.insert(0, "4")
        inner_pipe_amount_entry.insert(0, "2")
        is_multiple_pipes = True

    def initial_data_from_json():
        global initial_data_multiple_pipes, is_multiple_pipes
        inner_length = 0
        outer_length = 0
        inner_values = load_json_as_tuple('response_1721727588677.json')
        outer_values = load_json_as_tuple('response_1721727384287.json')

        with open('initial_data_json', 'w') as f:
            for i in inner_values:
                f.write(f"{to_float(i[1]['outer_diameter']['value'][:-2])/1000} ")
                f.write(f"{to_float(i[1]['inner_diameter']['value'][:-2])} ")
                f.write(f"{to_float(i[1]['length']['value'][:-2])}\n")
                inner_length += to_float(i[1]["length"]["value"][:-2])

            for i in outer_values:
                if not bool(i[1]["cased"]["value"]):
                    start, end = (i[1]["scope"]["value"]).split(" - ")
                    length = to_float(end[:-2]) - to_float((start[:-2]))
                    f.write(f"{length} ")
                    f.write(f"{to_float(i[1]['dolotos_diameter']['value'][:-2])/1000}\n")
                    outer_length += length
                else:
                    f.write(f"{to_float(i[1]['column_composition']['value'][0]['outer_diameter']['value'][:-2])/1000} ")
                    f.write(f"{to_float(i[1]['column_composition']['value'][0]['length']['value'][:-2])}\n")
                    outer_length += to_float(i[1]["column_composition"]["value"][0]["length"]["value"][:-2])

        with open('initial_data_json', 'r+') as f:
            content = f.read()
            f.seek(0)
            f.write(f"{inner_values[0][1]['number']} {outer_values[-1][1]['number']} {round(outer_length, 5)} {round(inner_length, 5)}\n" + content)

        initial_data_multiple_pipes = enter_initial_data_from_file_multiple_pipes("initial_data_json")
        save_values_message(initial_data_multiple_pipes)
        is_multiple_pipes = True

    content_frame.pack(fill="y", padx=30, pady=10, side="right")
    ctk.CTkLabel(content_frame, text="Автоматическая загрузка").pack(fill='x', pady=5)
    from_file_button_1 = ctk.CTkButton(content_frame, text="Загрузить данные (1) из WellPro", command=initial_data_from_file)
    from_file_button_2 = ctk.CTkButton(content_frame, text="Загрузить данные (2) из WellPro", command=initial_data_from_file_multiple_pipes)
    from_json_button = ctk.CTkButton(content_frame, text="Загрузить данные из json файлов", command=initial_data_from_json)
    from_file_button_1.pack(pady=10)
    from_file_button_2.pack(pady=10)
    from_json_button.pack(pady=10)
    ttk.Separator(content_frame, orient='horizontal').pack(fill='x', pady=10)
    label = ctk.CTkLabel(content_frame, text="Введите количество секций")
    label.pack(fill='x', pady=5)
    outer_pipe_amount_entry = create_entry(content_frame, "Введите количество элементов инструмента", "")
    inner_pipe_amount_entry = create_entry(content_frame, "Введите количество интервалов скважины", "")
    add_fields_button = ctk.CTkButton(content_frame, text="Добавить секции", command=add_pipe_fields)
    add_fields_button.pack(pady=10)

    '''
    ctk.CTkLabel(content_frame, text="Ввод данных для БТ").pack(fill='x')
    inner_diameter_entry_1 = ctk.CTkEntry(content_frame,
                                          placeholder_text="Внутренний диаметр БТ")
    inner_diameter_entry_1.pack(fill='x')
    outer_diameter_entry_1 = ctk.CTkEntry(content_frame,
                                          placeholder_text="Наружный диаметр БТ")
    outer_diameter_entry_1.pack(fill='x')
    section_length_entry_1 = ctk.CTkEntry(content_frame,
                                          placeholder_text="Длина секции БТ")
    section_length_entry_1.pack(fill='x')

    ctk.CTkLabel(content_frame, text="Ввод данных для ТБТ").pack(fill='x')
    inner_diameter_entry_2 = ctk.CTkEntry(content_frame,
                                          placeholder_text="Внутренний диаметр ТБТ")
    inner_diameter_entry_2.pack(fill='x')
    outer_diameter_entry_2 = ctk.CTkEntry(content_frame,
                                          placeholder_text="Наружный диаметр ТБТ")
    outer_diameter_entry_2.pack(fill='x')
    section_length_entry_2 = ctk.CTkEntry(content_frame,
                                          placeholder_text="Длина секции ТБТ")
    section_length_entry_2.pack(fill='x')

    ctk.CTkLabel(content_frame, text="Ввод данных для УБТ").pack(fill='x')
    inner_diameter_entry_3 = ctk.CTkEntry(content_frame,
                                          placeholder_text="Внутренний диаметр УБТ3")
    inner_diameter_entry_3.pack(fill='x')
    outer_diameter_entry_3 = ctk.CTkEntry(content_frame,
                                          placeholder_text="Наружный диаметр УБТ")
    outer_diameter_entry_3.pack(fill='x')
    section_length_entry_3 = ctk.CTkEntry(content_frame,
                                          placeholder_text="Длина секции УБТ")
    section_length_entry_3.pack(fill='x')

    ctk.CTkLabel(content_frame, text="Ввод данных для долота").pack(fill='x')
    inner_diameter_entry_4 = ctk.CTkEntry(content_frame,
                                          placeholder_text="Внутренний диаметр долота")
    inner_diameter_entry_4.pack(fill='x')
    outer_diameter_entry_4 = ctk.CTkEntry(content_frame,
                                          placeholder_text="Наружный диаметр долота")
    outer_diameter_entry_4.pack(fill='x')
    section_length_entry_4 = ctk.CTkEntry(content_frame,
                                          placeholder_text="Длина долота")
    section_length_entry_4.pack(fill='x')

    ctk.CTkLabel(content_frame, text="Ввод данных для обсаженного ствола").pack(fill='x')
    interval_diameter_entry_1 = ctk.CTkEntry(content_frame,
                                             placeholder_text="Диаметр интервала обсаженного ствола")
    interval_diameter_entry_1.pack(fill='x')
    interval_length_entry_1 = ctk.CTkEntry(content_frame,
                                           placeholder_text="Длина интервала обсаженного ствола")
    interval_length_entry_1.pack(fill='x')
    ctk.CTkLabel(content_frame, text="Ввод данных для открытого ствола").pack(fill='x')
    interval_diameter_entry_2 = ctk.CTkEntry(content_frame,
                                             placeholder_text="Диаметр интервала открытого ствола")
    interval_diameter_entry_2.pack(fill='x')
    interval_length_entry_2 = ctk.CTkEntry(content_frame,
                                           placeholder_text="Длина интервала открытого ствола")
    interval_length_entry_2.pack(fill='x')

    ctk.CTkLabel(content_frame, text="Ввод данных инструмента").pack(fill='x')
    current_depth_entry = ctk.CTkEntry(content_frame,
                                       placeholder_text="Глубина текущего забоя")
    current_depth_entry.pack(fill='x')
    lowering_depth_entry = ctk.CTkEntry(content_frame,
                                        placeholder_text="Глубина спуска инструмента")
    lowering_depth_entry.pack(fill='x')
    save_button = ctk.CTkButton(content_frame, text="Сохранить данные", command=save_multiple_pipes_data)
    save_button.pack(pady=10)
    '''


def user_data_input():
    def save_user_data():
        global user_data
        user_data = enter_user_data_manual(to_float(manifold_standpipe_pressure_drilling_entry.get()),
                                           to_float(pump_flow_rate_drilling_entry.get()),
                                           to_float(casing_pressure_test_entry.get()),
                                           to_float(formation_loss_pressure_entry.get()),
                                           to_float(weak_formation_fracture_pressure_entry.get()),
                                           to_float(drilling_mud_density_entry.get()),
                                           to_float(kill_fluid_viscosity_entry.get()),
                                           to_float(weak_formation_depth_entry.get()),
                                           to_float(standpipe_pressure_after_shut_in_entry.get()),
                                           to_float(wellhead_pressure_after_shut_in_entry.get()),
                                           to_float(additional_standpipe_pressure_during_injection_entry.get()),
                                           to_float(productive_zone_depth_entry.get()),
                                           to_float(fluid_influx_volume_entry.get()),
                                           to_float(lower_circulation_pressure_entry.get()),
                                           to_float(lower_consumption.get()))
        save_values_message(user_data)

    def user_data_from_file():
        global user_data, is_multiple_pipes
        if not is_multiple_pipes:
            user_data = enter_user_data_from_file("user_data")
        else:
            user_data = enter_user_data_from_file("user_data_multiple_pipes")
            manifold_standpipe_pressure_drilling_entry.insert(0, user_data["manifold_standpipe_pressure_drilling"])
            pump_flow_rate_drilling_entry.insert(0, user_data["pump_flow_rate_drilling"])
            casing_pressure_test_entry.insert(0, user_data["casing_pressure_test"])
            formation_loss_pressure_entry.insert(0, user_data["formation_loss_pressure"])
            weak_formation_fracture_pressure_entry.insert(0, user_data["weak_formation_fracture_pressure"])
            drilling_mud_density_entry.insert(0, user_data["drilling_mud_density"])
            kill_fluid_viscosity_entry.insert(0, user_data["kill_fluid_viscosity"])
            weak_formation_depth_entry.insert(0, user_data["weak_formation_depth"])
            standpipe_pressure_after_shut_in_entry.insert(0, user_data["standpipe_pressure_after_shut_in"])
            wellhead_pressure_after_shut_in_entry.insert(0, user_data["wellhead_pressure_after_shut_in"])
            additional_standpipe_pressure_during_injection_entry.insert(0, user_data["additional_standpipe_pressure_during_injection"])
            productive_zone_depth_entry.insert(0, user_data["productive_zone_depth"])
            fluid_influx_volume_entry.insert(0, user_data["fluid_influx_volume"])
            lower_circulation_pressure_entry.insert(0, user_data["lower_circulation_pressure"])
            lower_consumption.insert(0, user_data["lower_consumption"])
        save_values_message(user_data)

    content_frame.pack(fill="y", padx=30, pady=10, side="right")
    ctk.CTkLabel(content_frame, text="Автоматическая загрузка").pack(fill='x', pady=5)
    from_file_button = ctk.CTkButton(content_frame, text="Загрузить данные из WellPro", command=user_data_from_file)
    from_file_button.pack(pady=10)
    ttk.Separator(content_frame, orient='horizontal').pack(fill='x', pady=10)
    ctk.CTkLabel(content_frame, text="Исходные данные для расчета").pack(fill='x', pady=5)
    manifold_standpipe_pressure_drilling_entry = create_entry(content_frame, "Давление на стояке манифольда при бурении (промывке) скважины", "Па")
    pump_flow_rate_drilling_entry = create_entry(content_frame, "Подача насоса при бурении (промывке) скважины", "м³/с")
    casing_pressure_test_entry = create_entry(content_frame, "Давление опрессовки обсадной колонны", "Па")
    formation_loss_pressure_entry = create_entry(content_frame, "Давление начала поглощения слабого пласта", "Па")
    weak_formation_fracture_pressure_entry = create_entry(content_frame, "Давление гидроразрыва слабого пласта", "Па")
    drilling_mud_density_entry = create_entry(content_frame, "Плотность бурового раствора в скважине", "кг/м³")
    kill_fluid_viscosity_entry = create_entry(content_frame, "Вязкость жидкости глушения", "Па×с", entry_width=450)
    weak_formation_depth_entry = create_entry(content_frame, "Глубина залегания слабого пласта", "м", entry_width=450)
    standpipe_pressure_after_shut_in_entry = create_entry(content_frame, "Избыточное давление в бурильных трубах (на стояке) после закрытия скважины и стабилизации давления", "Па")
    wellhead_pressure_after_shut_in_entry = create_entry(content_frame, "Избыточное давление на устье скважины после её закрытия и стабилизации давления", "Па")
    additional_standpipe_pressure_during_injection_entry = create_entry(content_frame, "Дополнительное избыточное давление на стояке при закачке", "Па")
    productive_zone_depth_entry = create_entry(content_frame, "Глубина залегания продуктивного горизонта", "м")
    fluid_influx_volume_entry = create_entry(content_frame, "Объем поступившего в скважину флюида (увеличение уровня промывочной жидкости в рабочей емкости)", "м³")
    lower_circulation_pressure_entry = create_entry(content_frame, "Давление циркуляции при пониженной подаче насоса", "Па")
    lower_consumption = create_entry(content_frame, "Расход при пониженной подаче насоса", "м³/с")
    save_button = ctk.CTkButton(content_frame, text="Сохранить данные", command=save_user_data)
    save_button.pack(pady=10)

def add_well_data():
    def change_is_user_graph_base():
        global is_user_graph_base
        is_user_graph_base = not is_user_graph_base
        messagebox.showinfo("Сообщение", "Замеры на стояке добавлены")

    def change_is_user_graph_driller():
        global is_user_graph_driller
        is_user_graph_driller = not is_user_graph_driller
        messagebox.showinfo("Сообщение", "Замеры на устье добавлены")

    content_frame.pack(fill="y", padx=30, pady=10, side="right")
    ctk.CTkLabel(content_frame, text="Замеры давления").pack(fill='x', pady=5)
    user_graph_button = ctk.CTkButton(content_frame, text="Добавить замеры давления на стояке", command=change_is_user_graph_base)
    user_graph_button.pack(pady=10)
    user_graph_button = ctk.CTkButton(content_frame, text="Добавить замеры давления на устье", command=change_is_user_graph_driller)
    user_graph_button.pack(pady=10)


def start_method(method):
    global multiple_pipes_values, is_multiple_pipes, initial_data, initial_data_multiple_pipes
    if not is_multiple_pipes:
        drill_tool_internal_volume_value = drill_tool_internal_volume(initial_data['inner_diameter'],
                                                                      initial_data['section_length'])
        annular_volume_value = annular_volume(initial_data['interval_diameter'],
                                              initial_data['outer_diameter'],
                                              initial_data['section_length'],
                                              initial_data['current_depth'])
    else:
        initial_data = initial_data_multiple_pipes
        drill_tool_internal_volume_value = (
            calculate_drill_tool_internal_volume_multiple_pipes(
                [(val[2], val[1]) for val in initial_data_multiple_pipes['inner_pipe_values']]))
        volumes, lengths = (calculate_annular_volumes_multiple_pipes(
            [(val[1], val[0]) for val in initial_data_multiple_pipes['outer_pipe_values']],
            [(val[2], val[0]) for val in initial_data_multiple_pipes['inner_pipe_values']],
            initial_data_multiple_pipes['current_depth'] - initial_data_multiple_pipes['lowering_depth']))
        print(volumes, lengths)
        annular_volume_value = sum(volumes)
        lengths.insert(0, initial_data_multiple_pipes['current_depth'] - initial_data_multiple_pipes['lowering_depth'])
        annular_fluid_column_height_value = (
            calc_fluid_h(user_data['fluid_influx_volume'], [[i, j] for i, j in zip(volumes, lengths)]))
        #print("frfr", user_data['fluid_influx_volume'], initial_data_multiple_pipes['current_depth'], [[i, j] for i, j in zip(volumes, lengths)])
        #print("RAAR", [(val[1], val[0]) for val in initial_data_multiple_pipes['outer_pipe_values']],
        #      [(val[2], val[0]) for val in initial_data_multiple_pipes['inner_pipe_values']],
        #      initial_data_multiple_pipes['current_depth'] - initial_data_multiple_pipes['lowering_depth'])
        #print(initial_data_multiple_pipes['outer_pipe_values'])
        #print(initial_data_multiple_pipes['inner_pipe_values'])
        print([[i, j] for i, j in zip(volumes, lengths)])
        #print(initial_data_multiple_pipes['current_depth'], initial_data_multiple_pipes['lowering_depth'], drill_tool_internal_volume_value)
        #print(initial_data_multiple_pipes['current_depth'], [[i, j] for i, j in zip(volumes, lengths)])

        for i in range(0, 170, 1):
            print(i/10, end=" ")
            print(calc_water_h(i/10, 1, [[i, j] for i, j in zip(volumes, lengths)], initial_data_multiple_pipes['current_depth'], initial_data_multiple_pipes['lowering_depth'], initial_data_multiple_pipes['inner_pipe_values'], user_data['drilling_mud_density'], user_data['kill_fluid_viscosity'], volume_instr=drill_tool_internal_volume_value))
            print()
        print('\n\n\n\n\n\n')
        for i in range(0, 290, 1):
            print(i/10, end=" ")
            print(calc_water_h(i/10, 1, [[i, j] for i, j in zip(volumes, lengths)], initial_data_multiple_pipes['current_depth'], initial_data_multiple_pipes['lowering_depth'], initial_data_multiple_pipes['inner_pipe_values'], user_data['drilling_mud_density'], user_data['kill_fluid_viscosity'], outer_pipe_values=initial_data_multiple_pipes['outer_pipe_values']))
            print()

        #test_pr = calc_water_h(0, 20, [[i, j] for i, j in zip(volumes, lengths)], initial_data_multiple_pipes['current_depth'], initial_data_multiple_pipes['lowering_depth'], initial_data_multiple_pipes['inner_pipe_values'], drill_tool_internal_volume_value)
        #test2 = calc_water_h(10, 20, [[i, j] for i, j in zip(volumes, lengths)], initial_data_multiple_pipes['current_depth'], initial_data_multiple_pipes['lowering_depth'], initial_data_multiple_pipes['inner_pipe_values'])
        #test2_pr = calc_water_h(0.1, 20, [[i, j] for i, j in zip(volumes, lengths)], initial_data_multiple_pipes['current_depth'], initial_data_multiple_pipes['lowering_depth'], initial_data_multiple_pipes['inner_pipe_values'], drill_tool_internal_volume_value)
        #test3 = calc_water_h(100, 20, [[i, j] for i, j in zip(volumes, lengths)], initial_data_multiple_pipes['current_depth'], initial_data_multiple_pipes['lowering_depth'], initial_data_multiple_pipes['inner_pipe_values'])
        #test3_pr = calc_water_h(100, 20, [[i, j] for i, j in zip(volumes, lengths)], initial_data_multiple_pipes['current_depth'], initial_data_multiple_pipes['lowering_depth'], initial_data_multiple_pipes['inner_pipe_values'], drill_tool_internal_volume_value)
        '''
        print("t=1")
        print(test, test_pr)
        print("t=10")
        print(test2, test2_pr)
        print("t=100")
        print(test3, test3_pr)
        '''

    casing_max_allowable_pressure_value = casing_max_allowable_pressure(user_data['casing_pressure_test'])
    equivalent_density_beginning_absorption_value = equivalent_density_beginning_absorption(
        user_data['formation_loss_pressure'],
        user_data['weak_formation_depth'])
    equivalent_density_hydraulic_fracturing_value = equivalent_density_hydraulic_fracturing(
        user_data['weak_formation_fracture_pressure'],
        user_data['weak_formation_depth'])

    if not is_multiple_pipes:
        annular_space_volume_per_meter_value = annular_space_volume_per_meter(initial_data['interval_diameter'],
                                                                              initial_data['outer_diameter'])
        open_borehole_volume_per_meter_value = open_borehole_volume_per_meter(initial_data['interval_diameter'])
        annular_fluid_column_height_value = annular_fluid_column_height(user_data['fluid_influx_volume'],
                                                                        initial_data['current_depth'],
                                                                        initial_data['lowering_depth'],
                                                                        open_borehole_volume_per_meter_value,
                                                                        annular_space_volume_per_meter_value)
    reservoir_fluid_density_value = reservoir_fluid_density(user_data['wellhead_pressure_after_shut_in'],
                                                            user_data['standpipe_pressure_after_shut_in'],
                                                            annular_fluid_column_height_value)
    fluid_type_value = fluid_type(reservoir_fluid_density_value)
    calculate_reservoir_pressure_value = base_method.calculate_reservoir_pressure(
        user_data['standpipe_pressure_after_shut_in'],
        user_data['drilling_mud_density'],
        initial_data['current_depth'])
    calculation_density_killing_fluid_value = base_method.calculation_density_killing_fluid(
        calculate_reservoir_pressure_value,
        initial_data['current_depth'],
        user_data['productive_zone_depth'])
    if not is_multiple_pipes:
        calculate_initial_standpipe_pressure_value = base_method.calculate_initial_standpipe_pressure(
            user_data['pump_flow_rate_drilling'],
            initial_data['outer_diameter'],
            calculation_density_killing_fluid_value,
            user_data['kill_fluid_viscosity'],
            initial_data['section_length'],
            user_data['standpipe_pressure_after_shut_in'],
            user_data
            ['additional_standpipe_pressure_during_injection'])
        get_fluid_speed_value = base_method.get_fluid_speed(user_data['pump_flow_rate_drilling'],
                                                            initial_data['inner_diameter'])
        reynolds_criterion_value = base_method.reynolds_criterion(calculation_density_killing_fluid_value,
                                                                  get_fluid_speed_value,
                                                                  initial_data['inner_diameter'],
                                                                  user_data['kill_fluid_viscosity'])
        friction_coef_value = base_method.friction_coef(reynolds_criterion_value)
        get_tool_friction_loss_value = base_method.get_tool_friction_loss(friction_coef_value,
                                                                          calculation_density_killing_fluid_value,
                                                                          get_fluid_speed_value,
                                                                          initial_data['inner_diameter'],
                                                                          initial_data['section_length'])
        calculate_target_circulation_pressure_value = base_method.calculate_target_circulation_pressure(
            get_tool_friction_loss_value,
            user_data['additional_standpipe_pressure_during_injection'])
    else:
        values = initial_data['inner_pipe_values']
        length = len(values)
        get_fluid_speed_value = []
        reynolds_criterion_value = []
        friction_coef_value = []
        get_tool_friction_loss_value = 0
        for i in range(length):
            current_values = values[i]
            fluid_speed = base_method.get_fluid_speed(user_data['pump_flow_rate_drilling'], current_values[1])
            criterion = base_method.reynolds_criterion(calculation_density_killing_fluid_value, fluid_speed,
                                                       current_values[1], user_data['kill_fluid_viscosity'])
            friction_coef = base_method.friction_coef(criterion)
            friction_loss = base_method.get_tool_friction_loss(friction_coef, calculation_density_killing_fluid_value,
                                                               fluid_speed, current_values[1], current_values[2])

            get_fluid_speed_value.append(fluid_speed)
            reynolds_criterion_value.append(criterion)
            friction_coef_value.append(friction_coef)
            get_tool_friction_loss_value += friction_loss

        calculate_target_circulation_pressure_value = (base_method.calculate_target_circulation_pressure(
            get_tool_friction_loss_value, user_data['additional_standpipe_pressure_during_injection']))
        calculate_initial_standpipe_pressure_value = (calculate_target_circulation_pressure_value +
                                                      user_data['standpipe_pressure_after_shut_in'])

    global well_volumes, boundary_parameters, kick_analysis, base_method_values, driller_method_values
    well_volumes = {"Внутренний объём бурильного инструмента": round(drill_tool_internal_volume_value, 3),
                    "Объём кольцевого пространства": round(annular_volume_value, 3)}

    boundary_parameters = {"Максимально допустимое давление на устье": round(
        casing_max_allowable_pressure_value, 3),
        "Эквивалентная плотность начала поглощения": round(
            equivalent_density_beginning_absorption_value, 3),
        "Эквивалентная плотность начала гидроразрыва": round(
            equivalent_density_hydraulic_fracturing_value, 3)}
    _, boundary_parameters["Максимальное допускаемое давление на устье"] = [round(x, 3) for x in (
        driller_method.calc_max_wellhead_pressure(user_data["casing_pressure_test"],
                                                  user_data["formation_loss_pressure"],
                                                  user_data["drilling_mud_density"],
                                                  user_data["weak_formation_depth"]))]

    if not is_multiple_pipes:
        kick_analysis = {"Удельный объём 1 метра кольцевого пространства инструмент – скважина":
                             round(annular_space_volume_per_meter_value, 3),
                         "Удельный объём 1 метра открытого ствола": round(open_borehole_volume_per_meter_value, 3),
                         "Высота столба флюида в кольцевом пространстве (считая от забоя)":
                             round(annular_fluid_column_height_value, 3),
                         "Плотность пластового флюида": round(reservoir_fluid_density_value, 3),
                         "Тип флюида": fluid_type_value}

        base_method_values = {"Пластовое давление": round(calculate_reservoir_pressure_value, 3),
                              "Плотность жидкости глушения": round(calculation_density_killing_fluid_value, 3),
                              "Начальное давление на стояке": round(calculate_initial_standpipe_pressure_value, 3),
                              "Скорость потока в трубах": round(get_fluid_speed_value, 3),
                              "Критерий Рейнольдса в трубах": round(reynolds_criterion_value, 3),
                              "Коэффициент трения": round(friction_coef_value, 3),
                              "Потери давления на трение": round(get_tool_friction_loss_value, 3),
                              "Конечное давление на стояке": round(calculate_target_circulation_pressure_value, 3)}
    else:
        kick_analysis = {"Высота столба флюида в кольцевом пространстве (считая от забоя)":
                         round(annular_fluid_column_height_value, 3),
                         "Плотность пластового флюида": round(reservoir_fluid_density_value, 3),
                         "Тип флюида": fluid_type_value}

        base_method_values = {"Пластовое давление": round(calculate_reservoir_pressure_value, 3),
                              "Плотность жидкости глушения": round(calculation_density_killing_fluid_value, 3),
                              "Начальное давление на стояке": round(calculate_initial_standpipe_pressure_value, 3),
                              "Потери давления на трение": round(get_tool_friction_loss_value, 3),
                              "Конечное давление на стояке": round(calculate_target_circulation_pressure_value, 3)}
        length = len(initial_data['inner_pipe_values'])
        for i in range(length):
            base_method_values[f"Скорость потока в трубе {i+1}"] = round(get_fluid_speed_value[i], 3)
            base_method_values[f"Критерий Рейнольдса в трубе {i+1}"] = round(reynolds_criterion_value[i], 3)
            base_method_values[f"Коэффициент трения в трубе {i+1}"] = round(friction_coef_value[i], 3)
            #base_method_values[f"Коэффициент трения в трубе {i+1}"] = round(friction_coef_value[i], 3),

    if method == "Бурильщика":
        (boundary_parameters["Максимально допустимое давление на устье для недопущения поглощения"],
         boundary_parameters["Максимальное допускаемое давление на устье"]) = [round(x, 3) for x in (
            driller_method.calc_max_wellhead_pressure(user_data["casing_pressure_test"],
                                                      user_data["formation_loss_pressure"],
                                                      user_data["drilling_mud_density"],
                                                      user_data["weak_formation_depth"]))]
        calculate_pressure_circulation_value = driller_method.calculate_pressure_circulation(
            user_data["manifold_standpipe_pressure_drilling"],
            user_data["lower_consumption"],
            user_data["pump_flow_rate_drilling"],
            initial_data["lowering_depth"],
            initial_data["current_depth"])
        initial_circulation_pressure_value = driller_method.calc_initial_circulation_pressure(
            user_data["standpipe_pressure_after_shut_in"],
            calculate_pressure_circulation_value,
            user_data["additional_standpipe_pressure_during_injection"])
        calc_fact_reservoir_pressure_value = driller_method.calc_fact_reservoir_pressure(
            user_data["standpipe_pressure_after_shut_in"],
            user_data["drilling_mud_density"],
            user_data["productive_zone_depth"])
        calc_mud_density_increment_value = driller_method.calc_mud_density_increment(
            user_data["productive_zone_depth"],
            base_method_values["Пластовое давление"],
            user_data["drilling_mud_density"],
            initial_data["current_depth"])
        calc_kill_mud_density_value = driller_method.calc_kill_mud_density(
            base_method_values["Пластовое давление"],
            initial_data["current_depth"],
            user_data["productive_zone_depth"])
        calc_final_circulation_pressure_value = driller_method.calc_final_circulation_pressure(
            calculate_pressure_circulation_value,
            calc_kill_mud_density_value,
            user_data["drilling_mud_density"],
            user_data["additional_standpipe_pressure_during_injection"]
        )
        calc_gradient_value = driller_method.calc_gradient(
            user_data["standpipe_pressure_after_shut_in"],
            user_data["drilling_mud_density"],
            user_data["productive_zone_depth"]
        )
        calc_maximum_v_gas_value = driller_method.calc_maximum_v_gas(
            user_data["fluid_influx_volume"],
            annular_volume_value,
            calc_gradient_value
        )
        calc_maximum_pressure_annular_space_value = driller_method.calc_maximum_pressure_annular_space(
            base_method_values["Пластовое давление"],
            user_data["standpipe_pressure_after_shut_in"],
            calc_maximum_v_gas_value,
            annular_volume_value
        )
        calc_point_maximum_pressure_value = driller_method.calc_point_maximum_pressure(
            annular_volume_value,
            calc_maximum_v_gas_value
        )

        driller_method_values = {
            "Давление циркуляции на пониженной подаче": round(calculate_pressure_circulation_value, 3),
            "Начальное давление циркуляции": round(initial_circulation_pressure_value, 3),
            "Фактическое пластовое давление вскрытого пласта": round(calc_fact_reservoir_pressure_value, 3),
            "Дополнительное приращение плотности": round(calc_mud_density_increment_value, 3),
            "Плотность жидкости глушения": round(calc_kill_mud_density_value, 3),
            "Конечное давление циркуляции при глушении скважины": round(calc_final_circulation_pressure_value, 3),
            "Градиент распределения избыточного трубного давления": round(calc_gradient_value, 6),
            "Максимальное увеличение объема газовой пачки": round(calc_maximum_v_gas_value, 3),
            "Максимальное давление в кольцевом пространстве": round(calc_maximum_pressure_annular_space_value, 3),
            "Точка максимального давления на графике первой стадии глушения": round(
                calc_point_maximum_pressure_value, 3)
        }
    show_results(method=method)


def save_to_word(filename, *sections, fig=None, fig2=None):
    doc = Document()
    for section in sections:
        for title, data in section.items():
            doc.add_heading(title, level=1)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Название'
            hdr_cells[1].text = 'Значение'
            hdr_cells[2].text = 'Единица измерения'

            for key, value in data.items():
                row_cells = table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = str(value)
                row_cells[2].text = units.get(remove_end_num(key), "")

    if fig:
        graph_image_path = 'График.png'
        fig.savefig(graph_image_path)
        doc.add_heading('График', level=1)
        doc.add_picture(graph_image_path, width=Inches(6))

    if fig2:
        graph_image_path_2 = 'График2.png'
        fig2.savefig(graph_image_path_2)
        doc.add_heading('График 2', level=1)
        doc.add_picture(graph_image_path_2, width=Inches(6))

    doc.save(filename)


def user_plot(ax, file_path):
    values = {}
    with open(file_path, 'r') as file:
        for line in file:
            x, y = line.strip().split(' ')
            values[float(x)] = float(y)
    ax.plot(values.keys(), values.values(), color='black', label='Замеры со скважины')


def create_base_plot(standpipe_pressure_after_shut_in,
                     calculate_initial_standpipe_pressure_value,
                     drill_tool_internal_volume_value,
                     pump_flow_rate_drilling,
                     calculate_target_circulation_pressure_value,
                     annular_volume_value,
                     calc_max_wellhead_pressure_value, *,
                     is_driller=False,
                     lower_consumption=0,
                     initial_circulation_pressure=0,
                     calc_final_circulation_pressure=0
                     ):

    fig = Figure(figsize=(10, 6), dpi=100)
    ax = fig.add_subplot(111)

    ax.axhline(y=calc_max_wellhead_pressure_value, color='r', linestyle='--', label='Максимальное допускаемое давление')

    if is_driller:
        x1 = drill_tool_internal_volume_value / lower_consumption / 60
        x2 = annular_volume_value / lower_consumption / 60 + x1
        x3 = drill_tool_internal_volume_value / 60 / lower_consumption + x2
        x4 = annular_volume_value / 60 / lower_consumption + x3
        times1 = [0, 20, x1, x2, x3, x4]
        pressures1 = [standpipe_pressure_after_shut_in,
                      initial_circulation_pressure,
                      initial_circulation_pressure,
                      initial_circulation_pressure,
                      calc_final_circulation_pressure,
                      calc_final_circulation_pressure]

        ax.plot(times1, pressures1, color='y', label='Давление на стояке')
    else:
        times = [0,
                 2,
                 drill_tool_internal_volume_value / pump_flow_rate_drilling / 60,
                 (annular_volume_value / pump_flow_rate_drilling / 60) + (drill_tool_internal_volume_value / pump_flow_rate_drilling / 60)]
        pressures = [standpipe_pressure_after_shut_in,
                     calculate_initial_standpipe_pressure_value,
                     calculate_target_circulation_pressure_value,
                     calculate_target_circulation_pressure_value]

        ax.plot(times, pressures, marker='o', label='Давление на стояке')

    if is_user_graph_base:
        if is_driller:
            user_plot(ax, "user_graph_driller_1")
        else:
            user_plot(ax, "user_graph_base")

    ax.set_xlabel('Время (минуты)')
    ax.set_ylabel('Давление (МПа)')
    ax.set_title('Давление на стояке')
    ax.legend()

    return fig


def create_driller_plot(calc_max_wellhead_pressure_value,
                        drill_tool_internal_volume_value,
                        lower_consumption_value,
                        annular_volume_value,
                        standpipe_pressure_after_shut_in_value,
                        initial_circulation_pressure_value,
                        calc_final_circulation_pressure_value,
                        wellhead_pressure_after_shut_in_value,
                        additional_standpipe_pressure_during_injection_value,
                        calc_maximum_pressure_annular_space_value,
                        calc_point_maximum_pressure_value):

    fig = Figure(figsize=(10, 6), dpi=100)
    ax = fig.add_subplot(111)

    ax.axhline(y=calc_max_wellhead_pressure_value, color='r', linestyle='--',
               label='Максимальное избыточное давление на устье')

    x1 = drill_tool_internal_volume_value / lower_consumption_value / 60
    x2 = annular_volume_value / lower_consumption_value / 60 + x1
    x3 = drill_tool_internal_volume_value / 60 / lower_consumption_value + x2
    x4 = annular_volume_value / 60 / lower_consumption_value + x3

    times2 = [0, 20]
    pressures2 = [wellhead_pressure_after_shut_in_value,
                  wellhead_pressure_after_shut_in_value + additional_standpipe_pressure_during_injection_value]

    ax.plot(times2, pressures2, color='c', label='Рост затрубного давления в начале закачки')

    times3 = [x2, x3, x4]
    pressures3 = [standpipe_pressure_after_shut_in_value + additional_standpipe_pressure_during_injection_value,
                  standpipe_pressure_after_shut_in_value + additional_standpipe_pressure_during_injection_value,
                  additional_standpipe_pressure_during_injection_value]

    ax.plot(times3, pressures3, color='g', label='Затрубное давление во второй фазе')

    x_point = x1 + calc_point_maximum_pressure_value/lower_consumption_value/60
    times4 = [20, x_point]
    pressures4 = [wellhead_pressure_after_shut_in_value + additional_standpipe_pressure_during_injection_value,
                  calc_maximum_pressure_annular_space_value]
    x_values_1 = np.linspace(min(times4), max(times4), 100)
    y_values_1 = np.poly1d(np.polyfit(times4, pressures4, 2))(x_values_1)

    ax.plot(x_values_1, y_values_1, 'y-', label='Рост затрубного давления')

    times5 = [x_point, x2]
    pressures5 = [calc_maximum_pressure_annular_space_value,
                  standpipe_pressure_after_shut_in_value + additional_standpipe_pressure_during_injection_value]
    coefficients = np.polyfit(times5, pressures5, 2)

    def reflect_parabola(A, B, a, b, c):
        x1, y1 = A
        x2, y2 = B

        k = (y2 - y1) / (x2 - x1)
        d = y1 - k * x1

        new_a = -a
        new_b = 2 * k - b
        new_c = 2 * d - c

        return new_a, new_b, new_c

    coefficients = reflect_parabola((x_point, calc_maximum_pressure_annular_space_value),
                                    (x2, standpipe_pressure_after_shut_in_value + additional_standpipe_pressure_during_injection_value),
                                    *coefficients)
    x_values_2 = np.linspace(min(times5), max(times5), 100)
    y_values_2 = np.poly1d(coefficients)(x_values_2)

    ax.plot(x_values_2, y_values_2, 'b-', label='Снижение затрубного давления')

    if is_user_graph_driller:
        user_plot(ax, "user_graph_driller_2")

    ax.set_xlabel('Время (минуты)')
    ax.set_ylabel('Давление (МПа)')
    ax.set_title('Давление на устье')
    ax.legend()

    return fig


def show_results(method):
    def create_section(frame, title, data):
        section_label = ctk.CTkLabel(frame, text=title, font=('Arial', 16, 'bold'))
        section_label.pack(anchor='w', pady=(10, 0))

        table_frame = ctk.CTkFrame(frame)
        table_frame.pack(fill='x', pady=5)

        headers = ["Название", "       Значение", "Единица измерения"]
        header_widths = [470, 90, 40]

        for i, header in enumerate(headers):
            header_label = ctk.CTkLabel(table_frame, text=header, font=('Arial', 14, 'bold'), width=header_widths[i])
            header_label.grid(row=0, column=i, padx=5, pady=5, sticky="e")

        for row_index, (key, value) in enumerate(data.items(), start=1):
            unit = units.get(remove_end_num(key), "")
            key_label = ctk.CTkLabel(table_frame, text=key, anchor='w', wraplength=header_widths[0], width=header_widths[0])
            key_label.grid(row=row_index, column=0, padx=5, pady=5, sticky='w')
            value_label = ctk.CTkLabel(table_frame, text=value, anchor='e', wraplength=header_widths[1], width=header_widths[1])
            value_label.grid(row=row_index, column=1, padx=5, pady=5, sticky='e')
            unit_label = ctk.CTkLabel(table_frame, text=unit, anchor='e', wraplength=header_widths[2], width=header_widths[2])
            unit_label.grid(row=row_index, column=2, padx=5, pady=5, sticky='e')

        for col, width in enumerate(header_widths):
            table_frame.grid_columnconfigure(col, minsize=width, weight=1)

    content_frame.pack(fill="both", expand=True, padx=10, pady=10)

    create_section(content_frame, "Объёмы скважины", well_volumes)
    create_section(content_frame, "Параметры границ", boundary_parameters)
    create_section(content_frame, "Анализ ГНВП", kick_analysis)
    create_section(content_frame, "Параметры глушения", base_method_values)
    if method == "Бурильщика":
        create_section(content_frame, "Значения метода бурильщика", driller_method_values)
        base_fig = create_base_plot(user_data['standpipe_pressure_after_shut_in'],
                                    base_method_values["Начальное давление на стояке"],
                                    well_volumes["Внутренний объём бурильного инструмента"],
                                    user_data['pump_flow_rate_drilling'],
                                    base_method_values["Конечное давление на стояке"],
                                    well_volumes["Объём кольцевого пространства"],
                                    boundary_parameters["Максимальное допускаемое давление на устье"],
                                    is_driller=True,
                                    lower_consumption=user_data["lower_consumption"],
                                    initial_circulation_pressure=driller_method_values["Начальное давление циркуляции"],
                                    calc_final_circulation_pressure=driller_method_values["Конечное давление циркуляции при глушении скважины"])
    else:
        base_fig = create_base_plot(user_data['standpipe_pressure_after_shut_in'],
                                    base_method_values["Начальное давление на стояке"],
                                    well_volumes["Внутренний объём бурильного инструмента"],
                                    user_data['pump_flow_rate_drilling'],
                                    base_method_values["Конечное давление на стояке"],
                                    well_volumes["Объём кольцевого пространства"],
                                    boundary_parameters["Максимальное допускаемое давление на устье"])
    canvas = FigureCanvasTkAgg(base_fig, content_frame)
    canvas.draw()
    canvas.get_tk_widget().pack(fill='both', expand=True, padx=10, pady=10)
    if method == "Бурильщика":
        driller_fig = create_driller_plot(
            boundary_parameters["Максимально допустимое давление на устье для недопущения поглощения"],
            well_volumes["Внутренний объём бурильного инструмента"],
            user_data["lower_consumption"],
            well_volumes["Объём кольцевого пространства"],
            user_data['standpipe_pressure_after_shut_in'],
            driller_method_values["Начальное давление циркуляции"],
            driller_method_values["Конечное давление циркуляции при глушении скважины"],
            user_data['wellhead_pressure_after_shut_in'],
            user_data["additional_standpipe_pressure_during_injection"],
            driller_method_values["Максимальное давление в кольцевом пространстве"],
            driller_method_values["Точка максимального давления на графике первой стадии глушения"]
        )

        canvas = FigureCanvasTkAgg(driller_fig, content_frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill='both', expand=True, padx=10, pady=10)

    def choose_word_type():
        if method == "Базовый":
            save_to_word("Лист глушения (метод в один цикл).docx", {"Объёмы скважины": well_volumes},
                         {"Параметры границ": boundary_parameters}, {"Анализ ГНВП": kick_analysis},
                         {"Параметры глушения": base_method_values}, fig=base_fig)
        elif method == "Бурильщика":
            save_to_word("Лист глушения (метод бурильщика).docx", {"Объёмы скважины": well_volumes},
                         {"Параметры границ": boundary_parameters}, {"Анализ ГНВП": kick_analysis},
                         {"Параметры глушения": base_method_values}, {"Значения метода бурильщика": driller_method_values},
                         fig=base_fig, fig2=driller_fig)
        messagebox.showinfo("Информация", "Отчет скачан")

    ttk.Separator(content_frame, orient='horizontal').pack(fill='x', pady=10)
    from_file_button = ctk.CTkButton(content_frame, text="Скачать отчет", command=choose_word_type)
    from_file_button.pack(pady=10)


def show_initial_data():
    for widget in content_frame.winfo_children():
        widget.destroy()
    initial_data_input()


def show_initial_data_multiple_pipes():
    for widget in content_frame.winfo_children():
        widget.destroy()
    multiple_pipes_input()


def show_user_data():
    for widget in content_frame.winfo_children():
        widget.destroy()
    user_data_input()


def show_add_well_data():
    for widget in content_frame.winfo_children():
        widget.destroy()
    add_well_data()


def show_base_method_results():
    for widget in content_frame.winfo_children():
        widget.destroy()
    try:
        start_method(method="Базовый")
    except Exception as e:
        logging.error(traceback.format_exc())
        messagebox.showinfo("Ошибка", "Недостаточно данных для расчета базового метода", icon="error")


def show_driller_method_results():
    for widget in content_frame.winfo_children():
        widget.destroy()
    try:
        start_method(method="Бурильщика")
    except Exception as e:
        logging.error(traceback.format_exc())
        messagebox.showinfo("Ошибка", "Недостаточно данных для расчета метода бурильщика", icon="error")
    #ctk.CTkLabel(content_frame, text="Результаты метода бурильщика").pack(pady=10)
    #ctk.CTkButton(content_frame, text="Создать отчет", command=create_report).pack(pady=5)


def show_graphs():
    for widget in content_frame.winfo_children():
        widget.destroy()
    fig, ax = plt.subplots()
    times = [0, 2, 5, 10]
    pressures = [100, 200, 150, 300]
    ax.plot(times, pressures, 'r-')
    canvas = FigureCanvasTkAgg(fig, master=content_frame)
    canvas.draw()
    canvas.get_tk_widget().pack()


root = ctk.CTk()
root.title("Расчет параметров скважины")
background_image = ctk.CTkImage(Image.open("background.JPG"), size=(1920, 1080))
bg_label = ctk.CTkLabel(root, image=background_image, text="")
bg_label.place(x=0, y=0)

sidebar_frame = ctk.CTkFrame(root, corner_radius=5, width=500, fg_color="transparent")
sidebar_frame.pack(side='left', fill='both', expand=True, padx=10, pady=10)

buttons = [
    {"text": "Конструкция скважины", "command": show_initial_data_multiple_pipes},
    {"text": "Исходные данные для расчета", "command": show_user_data},
    {"text": "Данные со скважины", "command": show_add_well_data},
    {"text": "Результаты расчетов (метод в один цикл)", "command": show_base_method_results},
    {"text": "Результаты расчетов (метод бурильщика)", "command": show_driller_method_results},
    #{"text": "Графики", "command": show_graphs},
]

for btn in buttons:
    button = ctk.CTkButton(sidebar_frame, text=btn["text"], command=btn["command"], corner_radius=10, fg_color="white",
                           text_color="black")
    button.pack(pady=10, padx=10, fill='x')

content_frame = ctk.CTkScrollableFrame(root, corner_radius=5, width=frame_width, height=1080, fg_color="transparent",
                                       scrollbar_fg_color="transparent", scrollbar_button_color="white",
                                       scrollbar_button_hover_color="royal blue3")

ctk.set_appearance_mode("light")
ctk.set_default_color_theme("dark-blue")
root.after(0, lambda: root.state('zoomed'))

root.mainloop()
